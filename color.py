import json
import docx
import argparse
import sys


def index_fonts():
    palette=[]
    with open("palette_729.txt") as f:
        for line in f:
            line=line.strip()
            if not line.startswith("#"):
                continue
            palette.append(line)
    #print(palette)
    assert len(palette)==len(set(palette))

    rgb_colors=[]
    for colorspec in palette:
        color=docx.shared.RGBColor.from_string(colorspec[1:].upper())
        rgb_colors.append(color)

    return rgb_colors

rgb_colors = index_fonts()
print(f"{len(rgb_colors)} unique colors available in color palette")

def make_spans(color_map, paragraphs):
    all_results = [] # list of paragraphs where each is a list of (color, text)-tuples
    char_idx = -1
    for para in paragraphs:
        result = []
        last_color = None
        current_text = []
        for c in para:
            char_idx += 1
            color = color_map[char_idx]
            if last_color is None or last_color == color:
                current_text.append(c)
                last_color = color
            else:
                assert last_color is not None and len(current_text) > 0
                result.append((last_color, "".join(current_text)))
                last_color = color
                current_text = [c]
        else:
            assert last_color is not None and len(current_text) > 0
            result.append((last_color, "".join(current_text)))
        all_results.append(result)
    return all_results




def color_doc(doc_paragraphs, doc_idx, annotations, output_doc):
    
    color_map=[[] for _ in range(len("".join(p for p in doc_paragraphs)))] # this will store the span information for each char in text, init with empty

    # iterate through annotations
    for markable in annotations:
        # {"idx": "1", "text": "Introduction", "annotation": "1-abstract-new-cf1-1-sgl", "start": 0, "end": 12, "counter": 0}
        if "-sgl" in markable["annotation"] or markable["start"] is None or markable["end"] is None:
            #print("Skipping singleton markable:", markable)
            continue
        for lst in color_map[markable["start"]:markable["end"]]:
            lst.append(markable["idx"])

    # make a lookup table such that each unique overlap of answer ids has a color of its own whew
    color_lookup = {"": -1}
    for span in color_map: # span is a list of all markables which overlap as id strings
        span = "+".join(span) # make it a single string
        color_lookup.setdefault(span, len(color_lookup))

    #colors=index_fonts()

    spans=make_spans(color_map, doc_paragraphs)

    # write docx
    total_len=0
    for para_spans in spans:
        pgraph=output_doc.add_paragraph('')
        for color, txt in para_spans:
            r = pgraph.add_run(txt)
            font_idx = color_lookup.get("+".join(color), None)
            if font_idx >= 0:
                if font_idx > len(rgb_colors) - 1:
                    print(f"Warning, not enough colors, skipping {color} in document {doc_idx}!!")
                    continue
                r.font.color.rgb = rgb_colors[font_idx]
        total_len += len("".join(t for c, t in para_spans))+1

    return total_len, color_lookup


def main(args):

    with open(args.json, "rt", encoding="utf-8") as f:
        data = json.load(f)

    total_chars = 0
    file_counter = 0

    doc = docx.Document()
    metadata = []
    
    for i, d in enumerate(data):
    
        p = doc.add_paragraph()
        t = f"Document number {i}"
        total_chars += len(t) + 1
        p.add_run(t).bold=True
        
        p = doc.add_paragraph()
        t = f"Document identifier: {d['doc_id']}"
        total_chars += len(t) + 1
        p.add_run(t).bold=True
        
        l , color_map = color_doc(d["paragraphs"], d["doc_id"], d["annotations"], doc)
        total_chars += l
        metadata.append({"document number": i, "doc_id": d["doc_id"], "color_map": color_map})
        
        if total_chars > 900000: #document full!
            print(f"Total number of characters in this docx: {total_chars}")
            print(f"Saving to {args.output_dir}/file_{file_counter:03d}.docx")
            doc.save(f"{args.output_dir}/file_{file_counter:03d}.docx")
            file_counter += 1
            total_chars = 0
            doc = docx.Document()
            
    else:
        print(f"Total number of characters in this docx: {total_chars}")
        print(f"Saving to {args.output_dir}/file_{file_counter:03d}.docx")
        doc.save(f"{args.output_dir}/file_{file_counter:03d}.docx")
        


    with open(f"{args.output_dir}/meta.json", "wt", encoding="utf-8") as f:
        json.dump(metadata, f)
    






if __name__=="__main__":


    parser = argparse.ArgumentParser()
    parser.add_argument("--json", default=None, required=True, help="Input json file")
    parser.add_argument("--output_dir", default=None, required=True, help="Output docx directory")
    args = parser.parse_args()
    
    main(args)

